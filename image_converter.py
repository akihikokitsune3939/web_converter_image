"""
UNIVERSAL FILE CONVERTER
Supports: PNG, JPEG, WebP, BMP, GIF, TIFF, DOCX, PDF, TXT
"""

import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext
from PIL import Image, ImageTk
import os
import io
import sys
from datetime import datetime

# Optional imports with fallbacks
try:
    import fitz  # PyMuPDF
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False
    
try:
    from docx import Document
    from docx.shared import Inches, Pt
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

class UniversalConverter:
    def __init__(self, root):
        self.root = root
        self.root.title("Universal Converter Pro v2.0")
        self.root.geometry("900x700")
        
        # Variables
        self.file_path = None
        self.original_image = None
        self.preview_image = None
        self.file_type = None
        self.file_info = {}
        
        # Colors
        self.colors = {
            'bg': '#f5f5f5',
            'primary': '#2196F3',
            'success': '#4CAF50',
            'warning': '#FF9800',
            'danger': '#F44336',
            'dark': '#333333'
        }
        
        self.root.configure(bg=self.colors['bg'])
        self.setup_ui()
        self.center_window()
        
    def center_window(self):
        self.root.update_idletasks()
        width = self.root.winfo_width()
        height = self.root.winfo_height()
        x = (self.root.winfo_screenwidth() // 2) - (width // 2)
        y = (self.root.winfo_screenheight() // 2) - (height // 2)
        self.root.geometry(f'{width}x{height}+{x}+{y}')
    
    def setup_ui(self):
        # Main container
        main_container = tk.Frame(self.root, bg=self.colors['bg'])
        main_container.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
        
        # Header
        header_frame = tk.Frame(main_container, bg=self.colors['bg'])
        header_frame.pack(fill=tk.X, pady=(0, 20))
        
        tk.Label(header_frame, text="🔄 UNIVERSAL CONVERTER PRO", 
                font=("Arial", 22, "bold"), fg=self.colors['primary'],
                bg=self.colors['bg']).pack()
        
        tk.Label(header_frame, 
                text="PNG • JPEG • WebP • BMP • GIF • TIFF • DOCX • PDF • TXT", 
                font=("Arial", 10), fg=self.colors['dark'],
                bg=self.colors['bg']).pack()
        
        # Content area
        content_frame = tk.Frame(main_container, bg=self.colors['bg'])
        content_frame.pack(fill=tk.BOTH, expand=True)
        
        # Left panel - Preview
        left_panel = tk.LabelFrame(content_frame, text=" 📁 FILE PREVIEW", 
                                  font=("Arial", 11, "bold"),
                                  bg=self.colors['bg'], fg=self.colors['dark'],
                                  padx=15, pady=15)
        left_panel.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 10))
        
        # Preview area
        self.preview_canvas = tk.Canvas(left_panel, bg="white", 
                                       highlightthickness=1, 
                                       highlightbackground="#ddd")
        self.preview_canvas.pack(fill=tk.BOTH, expand=True)
        
        self.preview_text = scrolledtext.ScrolledText(left_panel, height=15,
                                                     font=("Consolas", 9),
                                                     wrap=tk.WORD)
        self.preview_text.pack_forget()  # Hidden by default
        
        self.preview_label = tk.Label(self.preview_canvas, 
                                     text="Select a file to preview",
                                     font=("Arial", 12), fg="gray",
                                     bg="white")
        self.preview_label.place(relx=0.5, rely=0.5, anchor=tk.CENTER)
        
        # Right panel - Controls
        right_panel = tk.Frame(content_frame, bg=self.colors['bg'])
        right_panel.pack(side=tk.RIGHT, fill=tk.Y)
        
        # File info panel
        info_frame = tk.LabelFrame(right_panel, text=" 📊 FILE INFORMATION", 
                                  font=("Arial", 11, "bold"),
                                  bg=self.colors['bg'], fg=self.colors['dark'],
                                  padx=15, pady=15)
        info_frame.pack(fill=tk.X, pady=(0, 15))
        
        self.info_display = tk.Text(info_frame, height=8, width=35,
                                   font=("Arial", 9), bg="#fafafa",
                                   relief=tk.FLAT)
        self.info_display.pack()
        self.info_display.config(state=tk.DISABLED)
        
        # Format selection
        format_frame = tk.LabelFrame(right_panel, text=" 🎯 CONVERT TO", 
                                    font=("Arial", 11, "bold"),
                                    bg=self.colors['bg'], fg=self.colors['dark'],
                                    padx=15, pady=15)
        format_frame.pack(fill=tk.X, pady=(0, 15))
        
        # Create scrollable format list
        format_container = tk.Frame(format_frame, bg=self.colors['bg'])
        format_container.pack(fill=tk.BOTH, expand=True)
        
        canvas = tk.Canvas(format_container, height=180, bg=self.colors['bg'],
                          highlightthickness=0)
        scrollbar = tk.Scrollbar(format_container, orient="vertical", 
                                command=canvas.yview)
        scrollable_frame = tk.Frame(canvas, bg=self.colors['bg'])
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        self.format_var = tk.StringVar(value="PNG")
        
        formats = [
            ("PNG Image (Lossless)", "PNG"),
            ("JPEG Image (Best for photos)", "JPEG"),
            ("WebP Image (Modern format)", "WebP"),
            ("BMP Image (Windows Bitmap)", "BMP"),
            ("GIF Image (Animation)", "GIF"),
            ("TIFF Image (High quality)", "TIFF"),
            ("PDF Document", "PDF"),
            ("DOCX Document (Word)", "DOCX"),
            ("TXT Text File", "TXT")
        ]
        
        for text, value in formats:
            rb = tk.Radiobutton(scrollable_frame, text=text, 
                               variable=self.format_var, value=value,
                               bg=self.colors['bg'], anchor="w",
                               font=("Arial", 9))
            rb.pack(fill=tk.X, pady=2, ipady=2)
        
        canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Buttons panel
        btn_frame = tk.Frame(right_panel, bg=self.colors['bg'])
        btn_frame.pack(fill=tk.X)
        
        buttons = [
            ("📂 SELECT FILE", self.select_file, self.colors['primary']),
            ("🔄 CONVERT FILE", self.convert_file, self.colors['success']),
            ("🗑️ CLEAR ALL", self.clear_all, self.colors['warning']),
            ("ℹ️ HELP", self.show_help, "#9C27B0"),
            ("🚪 EXIT", self.root.quit, self.colors['danger'])
        ]
        
        for text, command, color in buttons:
            btn = tk.Button(btn_frame, text=text, command=command,
                           bg=color, fg="white", font=("Arial", 10, "bold"),
                           padx=20, pady=10, cursor="hand2")
            btn.pack(fill=tk.X, pady=5)
            
            # Disable convert button initially
            if text == "🔄 CONVERT FILE":
                self.convert_btn = btn
                btn.config(state=tk.DISABLED)
    
    def select_file(self):
        filetypes = [
            ("All supported files", 
             "*.png *.jpg *.jpeg *.webp *.bmp *.gif *.tiff *.tif *.pdf *.docx *.txt"),
            ("Images", "*.png *.jpg *.jpeg *.webp *.bmp *.gif *.tiff *.tif"),
            ("Documents", "*.pdf *.docx *.txt"),
            ("All files", "*.*")
        ]
        
        path = filedialog.askopenfilename(filetypes=filetypes)
        if path:
            self.load_file(path)
    
    def load_file(self, path):
        self.file_path = path
        self.clear_preview()
        
        # Detect file type by extension
        ext = os.path.splitext(path)[1].lower()
        image_exts = ['.png', '.jpg', '.jpeg', '.webp', '.bmp', '.gif', '.tiff', '.tif']
        
        if ext in image_exts:
            self.file_type = 'image'
            self.load_image(path)
        elif ext == '.pdf':
            self.file_type = 'pdf'
            self.load_pdf(path)
        elif ext == '.docx':
            self.file_type = 'docx'
            self.load_docx(path)
        elif ext == '.txt':
            self.file_type = 'txt'
            self.load_txt(path)
        else:
            messagebox.showerror("Error", f"Unsupported file format: {ext}")
            return
        
        self.update_file_info()
        self.convert_btn.config(state=tk.NORMAL, bg=self.colors['success'])
    
    def load_image(self, path):
        try:
            self.original_image = Image.open(path)
            
            # Create preview
            canvas_width = self.preview_canvas.winfo_width() - 40
            canvas_height = self.preview_canvas.winfo_height() - 40
            
            if canvas_width < 10:  # If canvas not yet rendered
                canvas_width, canvas_height = 400, 300
            
            # Calculate resize ratio
            img_width, img_height = self.original_image.size
            ratio = min(canvas_width/img_width, canvas_height/img_height)
            new_size = (int(img_width * ratio), int(img_height * ratio))
            
            img_preview = self.original_image.copy()
            img_preview.thumbnail(new_size, Image.Resampling.LANCZOS)
            
            # Convert for display
            self.preview_image = ImageTk.PhotoImage(img_preview)
            
            # Clear and show image
            self.preview_canvas.delete("all")
            self.preview_canvas.create_image(canvas_width//2, canvas_height//2,
                                           image=self.preview_image,
                                           anchor=tk.CENTER)
            
            # Add overlay info
            info_text = f"{img_width}×{img_height} | {self.original_image.format}"
            self.preview_canvas.create_text(10, 10, text=info_text,
                                          anchor=tk.NW, fill="white",
                                          font=("Arial", 9, "bold"),
                                          fillbackground="black")
            
        except Exception as e:
            messagebox.showerror("Error", f"Cannot load image: {str(e)}")
    
    def load_pdf(self, path):
        if not PDF_SUPPORT:
            self.show_text_preview(f"PDF file: {os.path.basename(path)}\n\n"
                                  "PyMuPDF not installed.\n"
                                  "Install: pip install pymupdf")
            return
            
        try:
            pdf_doc = fitz.open(path)
            page = pdf_doc[0]
            
            # Render first page as image
            zoom = 2
            mat = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mat)
            img_data = pix.tobytes("ppm")
            
            img = Image.open(io.BytesIO(img_data))
            
            # Create preview
            canvas_width = self.preview_canvas.winfo_width() - 40
            canvas_height = self.preview_canvas.winfo_height() - 40
            
            img_width, img_height = img.size
            ratio = min(canvas_width/img_width, canvas_height/img_height)
            new_size = (int(img_width * ratio), int(img_height * ratio))
            
            img_preview = img.copy()
            img_preview.thumbnail(new_size, Image.Resampling.LANCZOS)
            
            self.preview_image = ImageTk.PhotoImage(img_preview)
            
            self.preview_canvas.delete("all")
            self.preview_canvas.create_image(canvas_width//2, canvas_height//2,
                                           image=self.preview_image,
                                           anchor=tk.CENTER)
            
            # Add PDF info
            info_text = f"PDF | Page 1/{len(pdf_doc)}"
            self.preview_canvas.create_text(10, 10, text=info_text,
                                          anchor=tk.NW, fill="white",
                                          font=("Arial", 9, "bold"),
                                          fillbackground="black")
            
            pdf_doc.close()
            
        except Exception as e:
            self.show_text_preview(f"PDF Content Preview\n\n"
                                  f"File: {os.path.basename(path)}\n"
                                  f"Error rendering: {str(e)}")
    
    def load_docx(self, path):
        if not DOCX_SUPPORT:
            self.show_text_preview(f"DOCX file: {os.path.basename(path)}\n\n"
                                  "python-docx not installed.\n"
                                  "Install: pip install python-docx")
            return
            
        try:
            doc = Document(path)
            
            # Extract text content
            text_content = ""
            for i, para in enumerate(doc.paragraphs[:20]):  # First 20 paragraphs
                if para.text.strip():
                    text_content += f"{para.text}\n"
            
            if len(doc.paragraphs) > 20:
                text_content += f"\n... and {len(doc.paragraphs) - 20} more paragraphs"
            
            self.show_text_preview(f"📝 DOCX CONTENT PREVIEW\n\n"
                                  f"File: {os.path.basename(path)}\n"
                                  f"Paragraphs: {len([p for p in doc.paragraphs if p.text.strip()])}\n"
                                  f"\n{'-'*40}\n\n"
                                  f"{text_content}")
            
        except Exception as e:
            self.show_text_preview(f"Error loading DOCX: {str(e)}")
    
    def load_txt(self, path):
        try:
            with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read(2000)  # Read first 2000 chars
                
            if len(content) == 2000:
                content += "\n\n... (truncated)"
            
            self.show_text_preview(f"📄 TEXT FILE PREVIEW\n\n"
                                  f"File: {os.path.basename(path)}\n"
                                  f"\n{'-'*40}\n\n"
                                  f"{content}")
            
        except Exception as e:
            self.show_text_preview(f"Error loading text file: {str(e)}")
    
    def show_text_preview(self, text):
        # Hide canvas, show text widget
        self.preview_canvas.pack_forget()
        self.preview_text.pack(fill=tk.BOTH, expand=True)
        
        self.preview_text.delete(1.0, tk.END)
        self.preview_text.insert(1.0, text)
        self.preview_text.config(state=tk.DISABLED)
    
    def show_image_preview(self):
        # Hide text widget, show canvas
        self.preview_text.pack_forget()
        self.preview_text.config(state=tk.NORMAL)
        self.preview_canvas.pack(fill=tk.BOTH, expand=True)
    
    def update_file_info(self):
        if not self.file_path:
            return
            
        self.info_display.config(state=tk.NORMAL)
        self.info_display.delete(1.0, tk.END)
        
        filename = os.path.basename(self.file_path)
        file_size = os.path.getsize(self.file_path)
        modified = datetime.fromtimestamp(os.path.getmtime(self.file_path))
        
        info = f"📄 File: {filename}\n"
        info += f"📁 Type: {self.file_type.upper()}\n"
        info += f"📂 Size: {self.format_size(file_size)}\n"
        info += f"📅 Modified: {modified.strftime('%Y-%m-%d %H:%M')}\n"
        info += f"📍 Path: {os.path.dirname(self.file_path)[:50]}...\n"
        info += "-" * 30 + "\n"
        
        # Format-specific info
        if self.file_type == 'image' and self.original_image:
            info += f"📐 Dimensions: {self.original_image.width}×{self.original_image.height}\n"
            info += f"🎨 Mode: {self.original_image.mode}\n"
            info += f"🔤 Format: {self.original_image.format}\n"
            
        elif self.file_type == 'pdf' and PDF_SUPPORT:
            try:
                pdf_doc = fitz.open(self.file_path)
                info += f"📑 Pages: {len(pdf_doc)}\n"
                info += f"📊 Version: {pdf_doc.pdf_version}\n"
                pdf_doc.close()
            except:
                info += "📑 PDF information unavailable\n"
                
        elif self.file_type == 'docx' and DOCX_SUPPORT:
            try:
                doc = Document(self.file_path)
                para_count = len([p for p in doc.paragraphs if p.text.strip()])
                info += f"📝 Paragraphs: {para_count}\n"
            except:
                info += "📝 DOCX information unavailable\n"
                
        elif self.file_type == 'txt':
            try:
                with open(self.file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    lines = sum(1 for _ in f)
                info += f"📝 Lines: {lines}\n"
            except:
                info += "📝 Text file information\n"
        
        self.info_display.insert(1.0, info)
        self.info_display.config(state=tk.DISABLED)
    
    def format_size(self, size_bytes):
        """Convert bytes to human readable format"""
        for unit in ['B', 'KB', 'MB', 'GB']:
            if size_bytes < 1024.0:
                return f"{size_bytes:.1f} {unit}"
            size_bytes /= 1024.0
        return f"{size_bytes:.1f} TB"
    
    def convert_file(self):
        if not self.file_path:
            messagebox.showwarning("Warning", "Please select a file first!")
            return
        
        target_format = self.format_var.get()
        
        # Check if conversion is possible
        if not self.validate_conversion(target_format):
            return
        
        # Get save path
        base_name = os.path.splitext(os.path.basename(self.file_path))[0]
        
        # Map formats to extensions
        format_ext = {
            'PNG': '.png', 'JPEG': '.jpg', 'WebP': '.webp',
            'BMP': '.bmp', 'GIF': '.gif', 'TIFF': '.tiff',
            'PDF': '.pdf', 'DOCX': '.docx', 'TXT': '.txt'
        }
        
        ext = format_ext.get(target_format, '')
        default_name = f"{base_name}_converted{ext}"
        
        # File type filters
        filetypes = [(f"{target_format} files", f"*{ext}"), ("All files", "*.*")]
        
        save_path = filedialog.asksaveasfilename(
            initialfile=default_name,
            defaultextension=ext,
            filetypes=filetypes
        )
        
        if not save_path:
            return  # User cancelled
        
        try:
            # Show progress
            self.show_progress("Converting...")
            
            # Perform conversion
            if target_format in ['PNG', 'JPEG', 'WebP', 'BMP', 'GIF', 'TIFF']:
                self.convert_to_image(save_path, target_format)
            elif target_format == 'PDF':
                self.convert_to_pdf(save_path)
            elif target_format == 'DOCX':
                self.convert_to_docx(save_path)
            elif target_format == 'TXT':
                self.convert_to_txt(save_path)
            
            self.hide_progress()
            
            # Show success message
            original_size = os.path.getsize(self.file_path)
            new_size = os.path.getsize(save_path)
            
            messagebox.showinfo("✅ Conversion Successful",
                              f"File converted successfully!\n\n"
                              f"Original: {self.format_size(original_size)}\n"
                              f"New: {self.format_size(new_size)}\n"
                              f"Saved as: {os.path.basename(save_path)}")
            
        except Exception as e:
            self.hide_progress()
            messagebox.showerror("Conversion Error",
                               f"Failed to convert file:\n\n{str(e)}")
    
    def validate_conversion(self, target_format):
        """Check if conversion from current format to target is possible"""
        image_formats = ['PNG', 'JPEG', 'WebP', 'BMP', 'GIF', 'TIFF']
        
        if target_format in image_formats:
            if self.file_type not in ['image', 'pdf']:
                messagebox.showwarning("Warning",
                                     f"Cannot convert {self.file_type.upper()} to image format.\n"
                                     f"Please select an image or PDF file.")
                return False
        
        return True
    
    def convert_to_image(self, save_path, target_format):
        """Convert to image format"""
        img = None
        
        # Get image from current file
        if self.file_type == 'image':
            img = self.original_image.copy()
        elif self.file_type == 'pdf' and PDF_SUPPORT:
            # Convert PDF first page to image
            pdf_doc = fitz.open(self.file_path)
            page = pdf_doc[0]
            pix = page.get_pixmap(matrix=fitz.Matrix(3, 3))
            img_data = pix.tobytes("ppm")
            img = Image.open(io.BytesIO(img_data))
            pdf_doc.close()
        else:
            raise ValueError(f"Cannot convert {self.file_type} to image")
        
        # Format-specific processing
        if target_format == 'JPEG':
            if img.mode in ('RGBA', 'LA', 'P'):
                # Add white background for transparency
                background = Image.new('RGB', img.size, (255, 255, 255))
                if img.mode == 'P':
                    img = img.convert('RGBA')
                background.paste(img, mask=img.split()[-1])
                img = background
            
            if img.mode != 'RGB':
                img = img.convert('RGB')
            
            img.save(save_path, 'JPEG', quality=95, optimize=True)
            
        elif target_format == 'PNG':
            img.save(save_path, 'PNG', compress_level=6)
            
        elif target_format == 'WebP':
            img.save(save_path, 'WEBP', quality=90, method=6)
            
        elif target_format == 'BMP':
            img.save(save_path, 'BMP')
            
        elif target_format == 'GIF':
            if img.mode not in ['P', 'L', 'RGB', 'RGBA']:
                img = img.convert('P', palette=Image.ADAPTIVE)
            img.save(save_path, 'GIF')
            
        elif target_format == 'TIFF':
            img.save(save_path, 'TIFF', compression='tiff_lzw')
    
    def convert_to_pdf(self, save_path):
        """Convert to PDF"""
        if self.file_type == 'image':
            img = self.original_image.copy()
            if img.mode == 'RGBA':
                img = img.convert('RGB')
            img.save(save_path, 'PDF', resolution=100.0)
            
        elif self.file_type == 'pdf':
            # PDF to PDF - just copy
            import shutil
            shutil.copy2(self.file_path, save_path)
            
        elif self.file_type == 'txt':
            # TXT to PDF
            if not PDF_SUPPORT:
                raise ImportError("PyMuPDF required for TXT to PDF conversion")
            
            pdf_doc = fitz.open()
            page = pdf_doc.new_page()
            
            with open(self.file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            # Simple text insertion
            page.insert_text((50, 50), content[:5000], fontsize=11)
            pdf_doc.save(save_path)
            pdf_doc.close()
            
        elif self.file_type == 'docx':
            # DOCX to PDF (requires both libraries)
            if not PDF_SUPPORT or not DOCX_SUPPORT:
                raise ImportError("Both PyMuPDF and python-docx required for DOCX to PDF")
            
            doc = Document(self.file_path)
            pdf_doc = fitz.open()
            page = pdf_doc.new_page()
            
            y = 50
            for para in doc.paragraphs:
                if para.text.strip():
                    page.insert_text((50, y), para.text, fontsize=12)
                    y += 20
                    if y > 750:  # New page if needed
                        page = pdf_doc.new_page()
                        y = 50
            
            pdf_doc.save(save_path)
            pdf_doc.close()
    
    def convert_to_docx(self, save_path):
        """Convert to DOCX"""
        if not DOCX_SUPPORT:
            raise ImportError("python-docx required for DOCX conversion")
        
        doc = Document()
        
        if self.file_type == 'image':
            doc.add_heading('Converted Image', 0)
            doc.add_paragraph(f"Original file: {os.path.basename(self.file_path)}")
            doc.add_paragraph(f"Converted on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            
            # Add image
            temp_img = "temp_conversion.jpg"
            self.original_image.save(temp_img, 'JPEG', quality=90)
            doc.add_picture(temp_img, width=Inches(5))
            os.remove(temp_img)
            
        elif self.file_type == 'txt':
            doc.add_heading('Converted Text File', 0)
            
            with open(self.file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            doc.add_paragraph(content)
            
        elif self.file_type == 'pdf':
            if not PDF_SUPPORT:
                raise ImportError("PyMuPDF required for PDF to DOCX")
            
            doc.add_heading('Converted PDF Document', 0)
            
            pdf_doc = fitz.open(self.file_path)
            for page_num in range(min(5, len(pdf_doc))):  # First 5 pages max
                page = pdf_doc[page_num]
                text = page.get_text()
                
                if text.strip():
                    doc.add_heading(f"Page {page_num + 1}", level=2)
                    doc.add_paragraph(text)
            
            pdf_doc.close()
            
        elif self.file_type == 'docx':
            # DOCX to DOCX - copy
            import shutil
            shutil.copy2(self.file_path, save_path)
            return
        
        doc.save(save_path)
    
    def convert_to_txt(self, save_path):
        """Convert to TXT"""
        content = ""
        
        if self.file_type == 'txt':
            # TXT to TXT - copy
            import shutil
            shutil.copy2(self.file_path, save_path)
            return
            
        elif self.file_type == 'docx' and DOCX_SUPPORT:
            doc = Document(self.file_path)
            for para in doc.paragraphs:
                if para.text.strip():
                    content += para.text + "\n\n"
                    
        elif self.file_type == 'pdf' and PDF_SUPPORT:
            pdf_doc = fitz.open(self.file_path)
            for page_num in range(len(pdf_doc)):
                page = pdf_doc[page_num]
                content += f"\n--- Page {page_num + 1} ---\n\n"
                content += page.get_text()
            pdf_doc.close()
            
        elif self.file_type == 'image':
            content = f"Image File: {os.path.basename(self.file_path)}\n"
            content += f"Dimensions: {self.original_image.width}×{self.original_image.height}\n"
            content += f"Format: {self.original_image.format}\n"
            content += f"Mode: {self.original_image.mode}\n"
            
        else:
            content = f"File: {os.path.basename(self.file_path)}\n"
            content += f"Type: {self.file_type}\n"
            content += "Content preview not available in text format.\n"
        
        # Write to file
        with open(save_path, 'w', encoding='utf-8') as f:
            f.write(content)
    
    def show_progress(self, message):
        """Show progress window"""
        self.progress = tk.Toplevel(self.root)
        self.progress.title("Please wait")
        self.progress.geometry("300x100")
        self.progress.transient(self.root)
        self.progress.grab_set()
        
        tk.Label(self.progress, text=message, font=("Arial", 11)).pack(pady=20)
        
        # Center progress window
        self.progress.update_idletasks()
        x = self.root.winfo_x() + (self.root.winfo_width() // 2) - (300 // 2)
        y = self.root.winfo_y() + (self.root.winfo_height() // 2) - (100 // 2)
        self.progress.geometry(f"+{x}+{y}")
    
    def hide_progress(self):
        """Hide progress window"""
        if hasattr(self, 'progress'):
            self.progress.destroy()
    
    def clear_preview(self):
        """Clear preview area"""
        self.show_image_preview()
        self.preview_canvas.delete("all")
        self.preview_label.place(relx=0.5, rely=0.5, anchor=tk.CENTER)
    
    def clear_all(self):
        """Clear everything"""
        self.file_path = None
        self.original_image = None
        self.preview_image = None
        self.file_type = None
        
        self.clear_preview()
        
        self.info_display.config(state=tk.NORMAL)
        self.info_display.delete(1.0, tk.END)
        self.info_display.config(state=tk.DISABLED)
        
        self.convert_btn.config(state=tk.DISABLED, bg="lightgray")
    
    def show_help(self):
        """Show help information"""
        help_text = """
        🔄 UNIVERSAL CONVERTER PRO - HELP
        
        SUPPORTED FORMATS:
        • Images: PNG, JPEG, WebP, BMP, GIF, TIFF
        • Documents: PDF, DOCX, TXT
        
        HOW TO USE:
        1. Click 'SELECT FILE' to choose a file
        2. Preview will appear automatically
        3. Select target format from the list
        4. Click 'CONVERT FILE'
        5. Choose save location
        
        INSTALLATION (if needed):
        • pip install pillow
        • pip install pymupdf (for PDF support)
        • pip install python-docx (for DOCX support)
        
        TIPS:
        • Large files may take longer to process
        • Some conversions require additional libraries
        • Keep original files as backups
        
        Need more help? Check the documentation.
        """
        
        help_window = tk.Toplevel(self.root)
        help_window.title("Help - Universal Converter")
        help_window.geometry("500x400")
        
        text_widget = scrolledtext.ScrolledText(help_window, wrap=tk.WORD)
        text_widget.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        text_widget.insert(1.0, help_text)
        text_widget.config(state=tk.DISABLED, font=("Arial", 10))
        
        # Center help window
        help_window.update_idletasks()
        x = self.root.winfo_x() + (self.root.winfo_width() // 2) - (500 // 2)
        y = self.root.winfo_y() + (self.root.winfo_height() // 2) - (400 // 2)
        help_window.geometry(f"+{x}+{y}")

def check_dependencies():
    """Check and report on required dependencies"""
    print("="*60)
    print("UNIVERSAL CONVERTER PRO - Dependency Check")
    print("="*60)
    
    # Required
    try:
        from PIL import Image
        print("✅ Pillow (required) - OK")
    except ImportError:
        print("❌ Pillow - MISSING (install: pip install pillow)")
    
    # Optional
    try:
        import fitz
        print("✅ PyMuPDF (optional) - OK")
    except ImportError:
        print("⚠️  PyMuPDF - NOT INSTALLED (PDF features limited)")
    
    try:
        from docx import Document
        print("✅ python-docx (optional) - OK")
    except ImportError:
        print("⚠️  python-docx - NOT INSTALLED (DOCX features limited)")
    
    print("="*60)
    print("Starting application...")

def main():
    check_dependencies()
    
    root = tk.Tk()
    app = UniversalConverter(root)
    
    # Set window icon if available
    try:
        root.iconbitmap('icon.ico')
    except:
        pass
    
    root.mainloop()

if __name__ == "__main__":
    main()